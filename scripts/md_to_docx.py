#!/usr/bin/env python3
"""Convert Markdown documentation to DOCX format with proper formatting."""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def setup_styles(doc):
    """Setup document styles for Chinese documentation."""
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Microsoft YaHei'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 128)
    h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Microsoft YaHei'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Microsoft YaHei'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Microsoft YaHei'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def parse_markdown(md_content):
    """Parse markdown content into structured elements."""
    elements = []
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Headers
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
        elif line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))
        elif line.startswith('##### '):
            elements.append(('h5', line[6:].strip()))
        
        # Horizontal rule
        elif line.strip() == '---':
            elements.append(('hr', ''))
        
        # Code block
        elif line.strip().startswith('```'):
            code_lines = []
            lang = line.strip()[3:].strip()
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append(('code', '\n'.join(code_lines), lang))
        
        # Table
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            elements.append(('table', table_lines))
            continue
        
        # Image
        elif line.strip().startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                alt, src = match.groups()
                elements.append(('image', alt, src))
        
        # Blockquote
        elif line.strip().startswith('>'):
            elements.append(('quote', line.strip()[1:].strip()))
        
        # List items
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            elements.append(('bullet', line.strip()[2:].strip()))
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.', '', line.strip()).strip()
            elements.append(('number', text))
        
        # Empty line
        elif not line.strip():
            elements.append(('empty', ''))
        
        # Regular paragraph
        else:
            elements.append(('paragraph', line.strip()))
        
        i += 1
    
    return elements


def add_formatted_paragraph(doc, text, style='Normal'):
    """Add paragraph with inline formatting (bold, italic, code)."""
    para = doc.add_paragraph(style=style)
    
    # Split by formatting markers
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        run = para.add_run()
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            run.text = part
        
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    return para


def create_docx(md_file, output_file, screenshots_dir=None):
    """Create DOCX from markdown file."""
    doc = Document()
    setup_styles(doc)
    
    # Read markdown content
    md_content = Path(md_file).read_text(encoding='utf-8')
    elements = parse_markdown(md_content)
    
    in_list = False
    list_items = []
    
    for elem in elements:
        elem_type = elem[0]
        
        if elem_type == 'h1':
            doc.add_heading(elem[1], level=1)
        elif elem_type == 'h2':
            doc.add_heading(elem[1], level=2)
        elif elem_type == 'h3':
            doc.add_heading(elem[1], level=3)
        elif elem_type == 'h4':
            doc.add_heading(elem[1], level=4)
        elif elem_type == 'h5':
            doc.add_heading(elem[1], level=5)
        
        elif elem_type == 'hr':
            doc.add_paragraph('_' * 50)
        
        elif elem_type == 'code':
            code_text = elem[1]
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.shading.background_pattern = True
        
        elif elem_type == 'table':
            table_lines = elem[1]
            if len(table_lines) >= 2:
                # Parse header
                headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = 'Microsoft YaHei'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                
                # Data rows (skip separator line)
                for line in table_lines[2:]:
                    cells = [c.strip() for c in line.split('|') if c.strip()]
                    if cells:
                        row_cells = table.add_row().cells
                        for i, cell_text in enumerate(cells[:len(headers)]):
                            row_cells[i].text = cell_text
                            for paragraph in row_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Microsoft YaHei'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        elif elem_type == 'image':
            alt_text = elem[1]
            img_path = elem[2]
            
            # Try to resolve image path
            if screenshots_dir:
                full_path = Path(screenshots_dir) / img_path.replace('./screenshots/', '').replace('/', '\\')
                if full_path.exists():
                    try:
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(str(full_path), width=Inches(5.5))
                        
                        # Add caption
                        caption = doc.add_paragraph(alt_text, style='Caption')
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in caption.runs:
                            run.font.size = Pt(9)
                            run.font.italic = True
                            run.font.name = 'Microsoft YaHei'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    except Exception as e:
                        doc.add_paragraph(f'[图片: {alt_text}]')
                else:
                    doc.add_paragraph(f'[图片: {alt_text} ({img_path})]')
            else:
                doc.add_paragraph(f'[图片: {alt_text}]')
        
        elif elem_type == 'quote':
            para = add_formatted_paragraph(doc, elem[1])
            para.paragraph_format.left_indent = Inches(0.3)
            for run in para.runs:
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.italic = True
        
        elif elem_type == 'bullet':
            add_formatted_paragraph(doc, elem[1], style='List Bullet')
        
        elif elem_type == 'number':
            add_formatted_paragraph(doc, elem[1], style='List Number')
        
        elif elem_type == 'paragraph':
            add_formatted_paragraph(doc, elem[1])
        
        elif elem_type == 'empty':
            pass  # Skip empty lines between paragraphs
    
    # Save document
    doc.save(output_file)
    print(f'DOCX saved to: {output_file}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python md_to_docx.py <input.md> <output.docx> [screenshots_dir]')
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_file = sys.argv[2]
    screenshots_dir = sys.argv[3] if len(sys.argv) > 3 else None
    
    create_docx(md_file, output_file, screenshots_dir)
